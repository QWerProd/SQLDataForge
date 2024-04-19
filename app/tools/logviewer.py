import wx
import os
import sqlite3

from docx import Document
from datetime import datetime
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wx.stc import StyledTextCtrl
from wx.lib.mixins.listctrl import ListCtrlAutoWidthMixin, ColumnSorterMixin
from app_parameters import APP_TEXT_LABELS, APP_PARAMETERS, APPLICATION_PATH
from app.error_catcher import ErrorCatcher


class Logviewer(wx.Frame):

    catcher = ErrorCatcher

    def get_errorlog(self) -> list:
        """Возвращает массив с массивами данных из t_error_log(error_code, text, date_catched)"""
        with sqlite3.connect(os.path.join(APPLICATION_PATH, 'app/app.db')) as app_conn:
            cursor = app_conn.cursor()
            try:
                log_data = cursor.execute(f"""SELECT el.error_code, lt.text, el.date_catched
                                              FROM t_error_log el
                                              JOIN t_lang_text lt ON el.error_code||'.CAPTION' = lt.label
                                              WHERE lt.lang = '{APP_PARAMETERS['APP_LANGUAGE']}'
                                              ORDER BY el.id DESC;""").fetchall()
            except sqlite3.Error as e:
                self.catcher.error_message('E014', 'sqlite_errorname: ' + e.sqlite_errorname + 'n' + str(e))
            finally:
                cursor.close()
        return log_data

    def get_execution_log(self) -> str:
        """Возвращает строку со всеми выполнеными запросами (первая запись - последний выполненный запрос)"""
        with sqlite3.connect(os.path.join(APPLICATION_PATH, 'app/app.db')) as app_conn:
            cursor = app_conn.cursor()
            try:
                log_data = cursor.execute(f"""SELECT el.query_text, el.date_execute
                                              FROM t_execution_log el
                                              ORDER BY el.id DESC;""").fetchall()
            except sqlite3.Error as e:
                self.catcher.error_message('E014', 'sqlite_errorname: ' + e.sqlite_errorname + 'n' + str(e))
            finally:
                cursor.close()

        log_text = ''
        for log_row in log_data:
            log_text += f'/*----------Query executed on {log_row[1]}----------*/\n' + log_row[0] + '\n\n'

        return log_text

    def generate_error_log_docx(self, event):
        """Генерация отчета DOCX на основе данных из get_error_log()"""
        path_to_file = ''
        with wx.DirDialog(None, APP_TEXT_LABELS['NEW_UDB_WIZARD.FIRST_PAGE.DIR_DIALOG'], ) as dirdlg:
            if dirdlg.ShowModal() == wx.ID_CANCEL:
                return
            path_to_file = dirdlg.GetPath()

        error_log_data = self.get_errorlog()

        # Составление документа
        doc = Document()
        sign_image = doc.add_picture(os.path.join(APPLICATION_PATH, 'img/SQLDataForge_sign.png'), width=Inches(2))
        doc.add_heading(APP_TEXT_LABELS['ERROR_LOG_FILE.DOCX.HEADING'], 0)

        # Добавление таблицы
        table = doc.add_table(rows = len(error_log_data) + 1, cols=3)
        table.style = 'Table Grid'
        table.cell(0, 0).text = APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CODE']
        table.cell(0, 1).text = APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CAPTION']
        table.cell(0, 2).text = APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CATCHED']
        for row in range(len(error_log_data)):
            for col in range(3):
                cell = table.cell(row+1, col)
                cell.text = error_log_data[row][col]

        doc.add_paragraph()
        date_time = doc.add_paragraph(str(datetime.now().strftime('%d %B %Y, %A, %X')))
        date_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        file_name = 'error_log_' + str(datetime.now().strftime('%d_%m_%Y_%H_%M_%S')) + '.docx'
        doc.save(os.path.join(path_to_file, file_name))

        return wx.MessageBox(APP_TEXT_LABELS['REPORT_WIZARD.REPORT_SAVED.MESSAGE'].format(file_name, path_to_file),
                             APP_TEXT_LABELS['REPORT_WIZARD.REPORT_SAVED.CAPTION'], wx.ICON_INFORMATION | wx.OK)

    def clear_error_log(self, event):
        """Очистка t_error_log"""
        with sqlite3.connect(os.path.join(APPLICATION_PATH, 'app/app.db')) as app_conn:
            cursor = app_conn.cursor()
            try:
                cursor.execute("DELETE FROM t_error_log;")
            except sqlite3.Error as e:
                catcher.error_message('E014', 'sqlite_errorname: ' + e.sqlite_errorname + 'n' + str(e))
            finally:
                cursor.close()

        self.errorlog_report.ClearLog()

    def generate_execution_log_docx(self, event):
        """Генерация отчета DOCX на основе данных из get_execution_log()"""
        path_to_file = ''
        with wx.DirDialog(None, APP_TEXT_LABELS['NEW_UDB_WIZARD.FIRST_PAGE.DIR_DIALOG'], ) as dirdlg:
            if dirdlg.ShowModal() == wx.ID_CANCEL:
                return
            path_to_file = dirdlg.GetPath()

        executionlog_data = self.get_execution_log()

        # Составление документа
        doc = Document()
        sign_image = doc.add_picture(os.path.join(APPLICATION_PATH, 'img/SQLDataForge_sign.png'), width=Inches(2))
        doc.add_heading(APP_TEXT_LABELS['EXECUTION_LOG_FILE.DOCX.HEADING'], 0)

        # Запись параграфов
        queries_pool = executionlog_data.split('\n\n')
        for query in queries_pool:
            code_run = doc.add_paragraph().add_run(query)
            code_run.font.size = Pt(8)
            code_run.font.name = 'Courier New'
            doc.add_paragraph()

        doc.add_paragraph()
        date_time = doc.add_paragraph(str(datetime.now().strftime('%d %B %Y, %A, %X')))
        date_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        file_name = 'execution_log_' + str(datetime.now().strftime('%d_%m_%Y_%H_%M_%S')) + '.docx'
        doc.save(os.path.join(path_to_file, file_name))

        return wx.MessageBox(APP_TEXT_LABELS['REPORT_WIZARD.REPORT_SAVED.MESSAGE'].format(file_name, path_to_file),
                             APP_TEXT_LABELS['REPORT_WIZARD.REPORT_SAVED.CAPTION'], wx.ICON_INFORMATION | wx.OK)

    def clear_execution_log(self, event):
        """Очистка t_execution_log"""
        with sqlite3.connect(os.path.join(APPLICATION_PATH, 'app/app.db')) as app_conn:
            cursor = app_conn.cursor()
            try:
                cursor.execute("DELETE FROM t_execution_log;")
            except sqlite3.Error as e:
                catcher.error_message('E014', 'sqlite_errorname: ' + e.sqlite_errorname + 'n' + str(e))
            finally:
                cursor.close()

        self.stc_redactor.ClearAll()

    def __init__(self, catcher: ErrorCatcher):
        wx.Frame.__init__(self, None, title=APP_TEXT_LABELS['MAIN.MAIN_MENU.TOOLS.LOGVIEWER'], size=(600, 450),
                          style=wx.CAPTION | wx.CLOSE_BOX | wx.MAXIMIZE_BOX | wx.MINIMIZE_BOX)
        self.SetMinSize((600, 450))
        self.SetIcon(wx.Icon(os.path.join(APPLICATION_PATH, 'img/main_icon.png'), wx.BITMAP_TYPE_PNG))
        self.catcher = catcher

        self.notebook = wx.Notebook(self, style=wx.NB_BOTTOM)

        # ----------

        self.errorlog_panel = wx.Panel(self.notebook)
        self.errorlog_sizer = wx.BoxSizer(wx.VERTICAL)
        self.errorlog_panel.SetSizer(self.errorlog_sizer)

        self.errorlog_report = ReportViewer(self.errorlog_panel,
                                            (APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CODE'],
                                             APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CAPTION'],
                                             APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.HEADER.ERROR_CATCHED']),
                                            self.get_errorlog())
        self.errorlog_sizer.Add(self.errorlog_report, 1, wx.EXPAND)

        # --------------------

        self.errorlog_buttons_panel = wx.Panel(self.errorlog_panel)
        self.errorlog_buttons_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.errorlog_buttons_panel.SetSizer(self.errorlog_buttons_sizer)

        self.erl_clear_button = wx.Button(self.errorlog_buttons_panel, label=APP_TEXT_LABELS['MAIN.MAIN_MENU.FILE.CLEAR_ALL'], size=(125, -1))
        self.erl_clear_button.Bind(wx.EVT_BUTTON, self.clear_error_log)
        self.erl_clear_button.Bind(wx.EVT_ENTER_WINDOW, lambda x: self.erl_clear_button.SetCursor(wx.Cursor(wx.CURSOR_HAND)))
        self.errorlog_buttons_sizer.Add(self.erl_clear_button, 0, wx.LEFT, 5)

        self.erl_save_docx_button = wx.Button(self.errorlog_buttons_panel, label=APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.BUTTONS.SAVE_DOCX'], size=(125, -1))
        self.erl_save_docx_button.Bind(wx.EVT_BUTTON, self.generate_error_log_docx)
        self.erl_save_docx_button.Bind(wx.EVT_ENTER_WINDOW, lambda x: self.erl_save_docx_button.SetCursor(wx.Cursor(wx.CURSOR_HAND)))
        self.errorlog_buttons_sizer.Add(self.erl_save_docx_button, 0, wx.LEFT, 5)

        self.errorlog_sizer.Add(self.errorlog_buttons_panel, 0, wx.ALIGN_RIGHT | wx.ALL, 5)
        # --------------------

        self.notebook.AddPage(self.errorlog_panel, APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG'])
        # ----------

        self.executionlog_panel = wx.Panel(self.notebook)
        self.executionlog_sizer = wx.BoxSizer(wx.VERTICAL)
        self.executionlog_panel.SetSizer(self.executionlog_sizer)

        self.stc_redactor = StyledTextCtrl(self.executionlog_panel)
        # Настройки шрифта
        self.stc_redactor.StyleSetFont(wx.stc.STC_STYLE_DEFAULT,
                                       wx.Font(pointSize=int(APP_PARAMETERS['STC_FONT_SIZE']),
                                               family=wx.FONTFAMILY_TELETYPE,
                                               style=wx.FONTSTYLE_NORMAL,
                                               weight=int(APP_PARAMETERS['STC_FONT_BOLD'])))
        self.stc_redactor.StyleClearAll()
        # Подсветка синтаксиса
        self.stc_redactor.SetLexer(wx.stc.STC_LEX_SQL)
        self.stc_redactor.SetKeyWords(0, APP_PARAMETERS['SQL_KEYWORDS'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_COMMENT, APP_PARAMETERS['STC_COLOUR_COMMENT'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_COMMENTLINE, APP_PARAMETERS['STC_COLOUR_COMMENT'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_COMMENTDOC, APP_PARAMETERS['STC_COLOUR_COMMENT'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_NUMBER, APP_PARAMETERS['STC_COLOUR_NUMBER'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_CHARACTER, APP_PARAMETERS['STC_COLOUR_STRING'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_STRING, APP_PARAMETERS['STC_COLOUR_STRING'])
        self.stc_redactor.StyleSetForeground(wx.stc.STC_SQL_WORD, APP_PARAMETERS['STC_COLOUR_WORD'])
        # Боковое поле
        self.stc_redactor.SetMarginType(1, wx.stc.STC_MARGIN_NUMBER)
        self.stc_redactor.SetMarginWidth(1, 45)

        self.stc_redactor.SetValue(self.get_execution_log())
        self.executionlog_sizer.Add(self.stc_redactor, 1, wx.EXPAND)

        # --------------------

        self.executionlog_buttons_panel = wx.Panel(self.executionlog_panel)
        self.executionlog_buttons_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.executionlog_buttons_panel.SetSizer(self.executionlog_buttons_sizer)

        self.exl_clear_button = wx.Button(self.executionlog_buttons_panel, label=APP_TEXT_LABELS['MAIN.MAIN_MENU.FILE.CLEAR_ALL'], size=(125, -1))
        self.exl_clear_button.Bind(wx.EVT_BUTTON, self.clear_execution_log)
        self.exl_clear_button.Bind(wx.EVT_ENTER_WINDOW, lambda x: self.exl_clear_button.SetCursor(wx.Cursor(wx.CURSOR_HAND)))
        self.executionlog_buttons_sizer.Add(self.exl_clear_button, 0, wx.LEFT, 5)

        self.exl_save_docx_button = wx.Button(self.executionlog_buttons_panel, label=APP_TEXT_LABELS['LOGVIEWER.ERROR_LOG.BUTTONS.SAVE_DOCX'], size=(125, -1))
        self.exl_save_docx_button.Bind(wx.EVT_BUTTON, self.generate_execution_log_docx)
        self.exl_save_docx_button.Bind(wx.EVT_ENTER_WINDOW, lambda x: self.exl_save_docx_button.SetCursor(wx.Cursor(wx.CURSOR_HAND)))
        self.executionlog_buttons_sizer.Add(self.exl_save_docx_button, 0, wx.LEFT, 5)

        self.executionlog_sizer.Add(self.executionlog_buttons_panel, 0, wx.ALL | wx.ALIGN_RIGHT, 5)
        # --------------------
        self.notebook.AddPage(self.executionlog_panel, APP_TEXT_LABELS['LOGVIEWER.EXECUTION_LOG'])
        # ----------


class ReportViewer(wx.ListCtrl, ListCtrlAutoWidthMixin, ColumnSorterMixin):
    columns = list
    choices = dict

    def __init__(self, parent: wx.Window, columns: list, choices: list):
        wx.ListCtrl.__init__(self, parent, style=wx.LC_REPORT | wx.LC_HRULES | wx.LC_VRULES)
        ListCtrlAutoWidthMixin.__init__(self)
        ColumnSorterMixin.__init__(self, len(columns))
        self.columns = columns
        self.choices = {}
        self.SetChoices(choices)

        # Appending headers
        self.SetColumns()
        self.SetItems()

    def SetColumns(self):
        for column in self.columns:
            self.AppendColumn(column)
            self.resizeColumn(100)
            self.SetHeaderAttr(wx.ItemAttr(self.GetForegroundColour(),
                                           self.GetBackgroundColour(),
                                           wx.Font(10, wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_BOLD)))

    def ClearLog(self):
        self.ClearAll()
        self.SetColumns()

    def SetChoices(self, choices: list):
        self.choices = {}
        for i in range(len(choices)):
            self.choices[i + 1] = choices[i]

    def SetItems(self):
        # Appending data
        self.itemDataMap = self.choices
        for key, value in self.choices.items():
            index = self.Append(value)
            self.SetItemData(index, key)

    def GetListCtrl(self):
        return self
