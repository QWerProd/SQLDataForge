import os
import re
import csv
import shutil
import sqlite3
import openpyxl
from docx import Document
from python_docx_replace import docx_replace

from sql_generator import SQLGenerator
from app_parameters import APPLICATION_PATH


class ReportGenerator:

    generator = SQLGenerator
    app_conn = sqlite3.Connection
    added_items = list
    rows_count = int
    file_name = str
    file_path = str
    report_info = dict
    template_name = str
    template_path = str

    def __init__(self, rows_count: int, added_items: list, file_name: str, file_path: str, report_info: dict,
                 template_name: str = None, template_path: str = None):
        self.added_items = added_items
        self.rows_count = rows_count
        self.file_name = file_name
        self.file_path = file_path
        self.report_info = report_info
        self.template_name = template_name
        self.template_path = template_path
        self.app_conn = sqlite3.connect(os.path.join(APPLICATION_PATH, 'app/app.db'))
        self.generator = SQLGenerator(self.app_conn, self.rows_count, self.added_items, None, is_simple_mode=False, is_format_columns=False)

    def generate_data(self):
        return self.generator.GenerateValues()

    def build_csv(self) -> bool:
        datadict = self.generate_data()
        with open(os.path.join(self.file_path, self.file_name), 'w', newline='') as csvfile:
            writer = csv.writer(csvfile, delimiter=';', quotechar='|', quoting=csv.QUOTE_MINIMAL)
            heading = []
            for keys in datadict.keys():
                heading.append(keys.split(':')[2])
            writer.writerow(heading)

            datarows = []
            for i in range(self.rows_count):
                temp_row = []
                for datarow in datadict.values():
                    temp_row.append(datarow[i])
                datarows.append(temp_row)
            writer.writerows(datarows)

        return True

    def build_docx(self) -> bool:
        datadict = self.generate_data()
        data = self.report_info
        for key, value in datadict.items():
            data[key.split(':')[2]] = value

        shutil.copy(self.template_path, os.path.join(self.file_path, self.file_name))
        template_file = Document(self.template_path)
        new_file = Document(os.path.join(self.file_path, self.file_name))

        replace_keys = {}

        # Записываем ключи для последующей замены
        for paragraph in new_file.paragraphs:
            for data_key in data.keys():
                if re.search(r'\$\{[\w-]*\}', paragraph.text):
                    data_value = data[data_key]
                    if isinstance(data_value, list):
                        data_value = data_value[0]
                    replace_keys[data_key] = data_value

        docx_replace(new_file, **replace_keys)

        new_file.save(os.path.join(self.file_path, self.file_name))
        return True

    def build_xlsx(self) -> bool:
        datadict = self.generate_data()
        data = self.report_info
        for key, value in datadict.items():
            data[key.split(':')[2]] = value

        shutil.copy(self.template_path, os.path.join(self.file_path, self.file_name))
        wordbook = openpyxl.load_workbook(os.path.join(self.file_path, self.file_name))
        sheet = wordbook.active

        row_num = 0
        null_rows_counter = 0
        for row in sheet.rows:
            is_null_row = True
            row_num += 1
            col_num = 0
            for cell in row:
                col_num += 1
                curr_value = cell.value
                is_contains_rep = True
                if re.search(r'\$\{[\w-]*\}', str(curr_value)) is not None:
                    is_null_row = False
                    while (1):
                        if re.search(r'\$\{TABLE\}', curr_value) is not None:
                            for col_key, col_values in data.items():
                                if col_key in ('report-name', 'report-date'):
                                    continue
                                table_row = row_num
                                col_width = 0
                                col_letter = sheet.cell(row=table_row, column=col_num).column_letter
                                for col_value in col_values:
                                    red_cell = sheet.cell(row=table_row, column=col_num)
                                    red_cell.value = col_value
                                    table_row += 1
                                    curr_col_width = len(str(red_cell.value))
                                    if col_width < curr_col_width:
                                        col_width = curr_col_width
                                col_num += 1
                                sheet.column_dimensions[col_letter].width = col_width + 5
                            is_contains_rep = False
                        else:
                            for data_key in self.report_info.keys():
                                if data_key in curr_value:
                                    cell.value = curr_value.replace(f'${{{data_key}}}', self.report_info[data_key])
                                    curr_value = cell.value
                            if re.search(r'\$\{[\w-]*\}', curr_value) is None:
                                is_contains_rep = False

                        if not is_contains_rep:
                            break
            if is_null_row:
                null_rows_counter += 1
            if null_rows_counter > 20:
                break

        wordbook.save(os.path.join(self.file_path, self.file_name))
        return True
